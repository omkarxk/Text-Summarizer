"""
File Handler Module
Handles reading and extracting text from various file formats
"""

import os
import tempfile
from typing import Tuple, Optional


class FileHandler:
    """
    Handles file operations and text extraction from various formats.
    Supports PDF, DOCX, DOC, and TXT files.
    """
    
    SUPPORTED_EXTENSIONS = ['pdf', 'doc', 'docx', 'txt']
    
    def __init__(self):
        """Initialize the file handler."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> dict:
        """
        Check which file processing libraries are available.
        
        Returns:
            Dictionary with availability status
        """
        self.available = {
            'pypdf2': False,
            'pdfplumber': False,
            'docx': False
        }
        
        try:
            import PyPDF2
            self.available['pypdf2'] = True
        except ImportError:
            pass
        
        try:
            import pdfplumber
            self.available['pdfplumber'] = True
        except ImportError:
            pass
        
        try:
            import docx
            self.available['docx'] = True
        except ImportError:
            pass
        
        return self.available
    
    def extract_from_pdf(self, file_path: str) -> Tuple[bool, str]:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (success, text or error message)
        """
        extracted_text = ''
        
        # Try PyPDF2 first
        if self.available['pypdf2']:
            try:
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + '\n'
                if extracted_text.strip():
                    return True, extracted_text.strip()
            except Exception as e:
                print(f"PyPDF2 error: {e}")
        
        # Try pdfplumber as fallback
        if self.available['pdfplumber']:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + '\n'
                if extracted_text.strip():
                    return True, extracted_text.strip()
            except Exception as e:
                print(f"pdfplumber error: {e}")
        
        if not self.available['pypdf2'] and not self.available['pdfplumber']:
            return False, "PDF processing libraries not installed. Please install PyPDF2 or pdfplumber."
        
        return False, "Could not extract text from PDF."
    
    def extract_from_docx(self, file_path: str) -> Tuple[bool, str]:
        """
        Extract text from a DOCX/DOC file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, text or error message)
        """
        if not self.available['docx']:
            return False, "python-docx library not installed. Please install it with: pip install python-docx"
        
        try:
            import docx
            doc = docx.Document(file_path)
            extracted_text = ''
            
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + '\n'
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text += cell.text + ' '
                    extracted_text += '\n'
            
            if extracted_text.strip():
                return True, extracted_text.strip()
            else:
                return False, "No text could be extracted from the document."
                
        except Exception as e:
            return False, f"Error reading document: {str(e)}"
    
    def extract_from_txt(self, file_path: str, encoding: str = 'utf-8') -> Tuple[bool, str]:
        """
        Extract text from a plain text file.
        
        Args:
            file_path: Path to the text file
            encoding: File encoding (default utf-8)
            
        Returns:
            Tuple of (success, text or error message)
        """
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            
            if text.strip():
                return True, text.strip()
            else:
                return False, "File is empty."
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return True, text.strip()
            except Exception as e:
                return False, f"Could not decode file: {str(e)}"
        except Exception as e:
            return False, f"Error reading file: {str(e)}"
    
    def extract_text(self, file_path: str) -> Tuple[bool, str]:
        """
        Extract text from any supported file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (success, text or error message)
        """
        if not os.path.exists(file_path):
            return False, "File not found."
        
        # Get file extension
        extension = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
        
        if extension == 'pdf':
            return self.extract_from_pdf(file_path)
        elif extension in ['doc', 'docx']:
            return self.extract_from_docx(file_path)
        elif extension == 'txt':
            return self.extract_from_txt(file_path)
        else:
            # Try to read as plain text
            return self.extract_from_txt(file_path)
    
    def save_temp_file(self, file_content: bytes, extension: str) -> str:
        """
        Save uploaded file content to a temporary file.
        
        Args:
            file_content: Raw file bytes
            extension: File extension
            
        Returns:
            Path to temporary file
        """
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{extension}') as tmp:
            tmp.write(file_content)
            return tmp.name
    
    def cleanup_temp_file(self, file_path: str) -> bool:
        """
        Remove a temporary file.
        
        Args:
            file_path: Path to file to remove
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            return True
        except Exception:
            return False
    
    def get_supported_formats(self) -> dict:
        """
        Get information about supported file formats.
        
        Returns:
            Dictionary with format information and availability
        """
        return {
            'pdf': {
                'supported': self.available['pypdf2'] or self.available['pdfplumber'],
                'libraries': ['PyPDF2', 'pdfplumber']
            },
            'docx': {
                'supported': self.available['docx'],
                'libraries': ['python-docx']
            },
            'doc': {
                'supported': self.available['docx'],
                'libraries': ['python-docx']
            },
            'txt': {
                'supported': True,
                'libraries': ['built-in']
            }
        }
